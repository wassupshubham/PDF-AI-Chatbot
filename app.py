import os
import sys

# Reconfigure console streams to use UTF-8 to prevent charmap UnicodeEncodeErrors on Windows
if hasattr(sys.stdout, 'reconfigure'):
    sys.stdout.reconfigure(encoding='utf-8')
if hasattr(sys.stderr, 'reconfigure'):
    sys.stderr.reconfigure(encoding='utf-8')

# 1. Core Torch Imports right on top to resolve any upstream library type-hint bugs
import torch
import torch.nn as nn
sys.modules['nn'] = nn

import streamlit as st
from dotenv import load_dotenv
import io

# Load environment variables
load_dotenv()

# --- CORE RAG PIPELINE FUNCTIONS ---
import pdfplumber
from pypdf import PdfReader
import easyocr
from PIL import Image

@st.cache_resource
def load_local_ocr():
    return easyocr.Reader(['en'], gpu=False)

import pypdfium2 as pdfium

def extract_text_from_pdf(pdf_file):
    text = ""
    reader = load_local_ocr()
    
    try:
        with pdfplumber.open(pdf_file) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text(layout=True)
                if page_text:
                    text += page_text + "\n"
        
        if not text.strip() or len(text.strip()) < 100:
            with st.spinner("🧠 Scanned Layout Detected. Initializing Visual OCR Engine..."):
                pdf_file.seek(0)
                pdf_data = pdf_file.read()
                pdf_doc = pdfium.PdfDocument(pdf_data)
                
                for page_num, page in enumerate(pdf_doc):
                    bitmap = page.render(scale=2)
                    pil_img = bitmap.to_pil()
                    results = reader.readtext(pil_img, detail=0)
                    if results:
                        text += f"\n--- Page {page_num+1} Scan ---\n" + " ".join(results) + "\n"
                                    
    except Exception as e:
        st.error(f"Local Extraction Layer Warning: {e}")
    return text

import docx

def extract_text_from_docx(docx_file):
    try:
        doc = docx.Document(docx_file)
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
        return '\n'.join(full_text)
    except Exception as e:
        st.error(f"Word Ingestion Error: {e}")
        return ""

def extract_text_from_txt(txt_file):
    try:
        txt_file.seek(0)
        return txt_file.read().decode("utf-8", errors="ignore")
    except Exception as e:
        st.error(f"Text Ingestion Error: {e}")
        return ""

def extract_text_from_file(uploaded_file):
    file_name = uploaded_file.name.lower()
    if file_name.endswith('.pdf'):
        return extract_text_from_pdf(uploaded_file)
    elif file_name.endswith('.docx'):
        return extract_text_from_docx(uploaded_file)
    elif file_name.endswith(('.txt', '.md')):
        return extract_text_from_txt(uploaded_file)
    else:
        st.error("❌ Unsupported file format detected.")
        return ""

from langchain_text_splitters import RecursiveCharacterTextSplitter
def split_text(text):
    return RecursiveCharacterTextSplitter(chunk_size=600, chunk_overlap=150).split_text(text)

from langchain_huggingface import HuggingFaceEmbeddings
def get_embeddings():
    return HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")

# 🔥 FIX: Using an explicit Ephemeral Client (Pure RAM In-Memory) to bypass HuggingFace's disk write protection
import chromadb
from langchain_chroma import Chroma

def create_vectorstore(chunks, embeddings):
    # EphemeralClient memory space mein initialize hota hai aur disk par touch nahi karta
    chroma_client = chromadb.EphemeralClient()
    
    return Chroma.from_texts(
        texts=chunks, 
        embedding=embeddings,
        client=chroma_client
    )

from langchain_google_genai import ChatGoogleGenerativeAI
from langchain_groq import ChatGroq

def get_llm(provider="Gemini", model_name=None, temperature=0.0):
    if provider == "Gemini":
        model = model_name or "gemini-2.5-flash"
        return ChatGoogleGenerativeAI(
            model=model,
            google_api_key=os.getenv("GOOGLE_API_KEY"),
            temperature=temperature
        )
    else:
        model = model_name or "llama-3.1-8b-instant"
        return ChatGroq(
            model=model,
            groq_api_key=os.getenv("GROQ_API_KEY"),
            temperature=temperature
        )


# --- SMARTPHONE OPTIMIZED UI INTERFACE ---
st.set_page_config(page_title="InsightDocs AI", page_icon="🤖", layout="wide")

st.markdown("""
    <style>
    @import url('https://fonts.googleapis.com/css2?family=Outfit:wght@300;400;500;600;700;800&family=Playfair+Display:ital,wght@0,600;1,600&display=swap');

    /* Global Reset and Obsidian Dark Background */
    .stApp {
        background: radial-gradient(circle at top left, #0D0F16, #07090D 70%) !important;
        color: #E5E7EB !important;
        font-family: 'Outfit', sans-serif !important;
    }

    /* Hide Default Headers and Footers */
    [data-testid="stHeader"], footer, #MainMenu {
        display: none !important;
    }

    /* Sidebar Glassmorphic Styling */
    section[data-testid="stSidebar"] {
        background: rgba(10, 12, 20, 0.6) !important;
        backdrop-filter: blur(20px) !important;
        -webkit-backdrop-filter: blur(20px) !important;
        border-right: 1px solid rgba(255, 255, 255, 0.05) !important;
    }

    /* Custom Settings Typography in Sidebar */
    section[data-testid="stSidebar"] h3 {
        color: #F3F4F6 !important;
        font-weight: 700 !important;
        letter-spacing: -0.5px !important;
        font-family: 'Outfit', sans-serif !important;
    }

    /* Title & Subtitle Styling with Flowing Animation */
    .breathing-header {
        font-family: 'Outfit', sans-serif !important;
        font-size: 2.8rem !important;
        font-weight: 800 !important;
        text-align: center !important;
        margin-top: 15px !important;
        margin-bottom: 5px !important;
        background: linear-gradient(135deg, #60A5FA 0%, #A78BFA 50%, #EC4899 100%) !important;
        background-size: 200% auto !important;
        -webkit-background-clip: text !important;
        -webkit-text-fill-color: transparent !important;
        animation: shine 6s linear infinite, breathe 4s ease-in-out infinite !important;
    }

    .sub-title {
        font-family: 'Outfit', sans-serif !important;
        font-size: 1.1rem !important;
        color: #9CA3AF !important;
        text-align: center !important;
        margin-bottom: 30px !important;
        font-weight: 300 !important;
        letter-spacing: 0.5px !important;
    }

    /* Animations */
    @keyframes shine {
        to { background-position: 200% center; }
    }

    @keyframes breathe {
        0%, 100% {
            transform: scale(1);
            filter: drop-shadow(0 0 8px rgba(167, 139, 250, 0.25));
        }
        50% {
            transform: scale(1.015);
            filter: drop-shadow(0 0 20px rgba(167, 139, 250, 0.55));
        }
    }

    @keyframes pulseGlow {
        0%, 100% {
            box-shadow: 0 4px 15px rgba(139, 92, 246, 0.15), inset 0 1px 0 rgba(255, 255, 255, 0.05);
        }
        50% {
            box-shadow: 0 4px 30px rgba(139, 92, 246, 0.35), inset 0 1px 0 rgba(255, 255, 255, 0.08);
        }
    }

    /* AI Response Box: Obsidian dark background with glowing color cycle gradient border */
    .ai-response-box {
        background: rgba(11, 14, 23, 0.8) !important;
        backdrop-filter: blur(12px) !important;
        border-radius: 14px !important;
        border: 1px solid rgba(255, 255, 255, 0.05) !important;
        border-left: 5px solid #8B5CF6 !important;
        padding: 24px !important;
        color: #F3F4F6 !important;
        line-height: 1.7 !important;
        margin-top: 15px !important;
        font-size: 1.05rem !important;
        box-shadow: 0 4px 25px rgba(0, 0, 0, 0.25) !important;
        animation: border-cycle 6s linear infinite, pulseGlow 4s ease-in-out infinite !important;
    }

    @keyframes border-cycle {
        0%, 100% { border-left-color: #3B82F6; }
        33% { border-left-color: #8B5CF6; }
        66% { border-left-color: #EC4899; }
    }

    /* File Uploader Style Customization */
    [data-testid="stFileUploader"] {
        background: rgba(22, 28, 45, 0.2) !important;
        border: 1px dashed rgba(255, 255, 255, 0.1) !important;
        border-radius: 12px !important;
        padding: 15px !important;
        transition: all 0.3s cubic-bezier(0.4, 0, 0.2, 1) !important;
    }

    [data-testid="stFileUploader"]:hover {
        border-color: #8B5CF6 !important;
        background: rgba(22, 28, 45, 0.35) !important;
        box-shadow: 0 0 15px rgba(139, 92, 246, 0.1) !important;
    }

    /* Custom Interactive Buttons styling */
    div.stButton > button {
        background: linear-gradient(135deg, #3B82F6 0%, #8B5CF6 50%, #EC4899 100%) !important;
        background-size: 200% auto !important;
        color: white !important;
        border: none !important;
        border-radius: 10px !important;
        height: 48px !important;
        font-weight: 700 !important;
        letter-spacing: 0.5px !important;
        transition: all 0.4s cubic-bezier(0.4, 0, 0.2, 1) !important;
        box-shadow: 0 4px 15px rgba(139, 92, 246, 0.2) !important;
        animation: button-pulse 3s infinite !important;
    }

    div.stButton > button:hover {
        background-position: right center !important;
        transform: translateY(-2px) scale(1.01) !important;
        box-shadow: 0 6px 22px rgba(139, 92, 246, 0.45) !important;
    }

    @keyframes button-pulse {
        0%, 100% { box-shadow: 0 4px 15px rgba(139, 92, 246, 0.2); }
        50% { box-shadow: 0 4px 25px rgba(139, 92, 246, 0.45); }
    }

    /* Text Input custom styles */
    div[data-baseweb="input"] {
        background-color: rgba(13, 16, 27, 0.7) !important;
        border: 1px solid rgba(255, 255, 255, 0.08) !important;
        border-radius: 10px !important;
        padding: 2px 5px !important;
        transition: all 0.3s ease !important;
    }

    div[data-baseweb="input"]:focus-within {
        border-color: #8B5CF6 !important;
        box-shadow: 0 0 12px rgba(139, 92, 246, 0.25) !important;
        background-color: rgba(13, 16, 27, 0.9) !important;
    }

    /* Custom Status / Success Badges styling */
    div[data-testid="stNotification"] {
        background: rgba(16, 24, 40, 0.6) !important;
        backdrop-filter: blur(10px) !important;
        border-radius: 10px !important;
        border: 1px solid rgba(255, 255, 255, 0.05) !important;
        box-shadow: 0 4px 15px rgba(0, 0, 0, 0.15) !important;
    }

    /* Custom Chat Message Styling */
    div[data-testid="stChatMessage"] {
        background: rgba(22, 28, 45, 0.25) !important;
        border: 1px solid rgba(255, 255, 255, 0.05) !important;
        border-radius: 14px !important;
        padding: 18px !important;
        box-shadow: 0 4px 15px rgba(0, 0, 0, 0.15) !important;
        margin-bottom: 15px !important;
    }
    div[data-testid="stChatMessage"] [data-testid="stChatMessageContent"] {
        color: #F3F4F6 !important;
    }
    
    /* Custom Chat Input styling */
    div[data-testid="stChatInput"] {
        background-color: transparent !important;
        border: none !important;
        margin-top: 15px !important;
    }
    div[data-testid="stChatInput"] textarea {
        background-color: rgba(13, 16, 27, 0.85) !important;
        border: 1px solid rgba(255, 255, 255, 0.08) !important;
        color: #F3F4F6 !important;
        border-radius: 12px !important;
        padding: 12px !important;
        box-shadow: 0 4px 25px rgba(0, 0, 0, 0.3) !important;
        transition: all 0.3s ease !important;
    }
    div[data-testid="stChatInput"] textarea:focus {
        border-color: #8B5CF6 !important;
        box-shadow: 0 0 18px rgba(139, 92, 246, 0.35) !important;
    }
    </style>
""", unsafe_allow_html=True)

# --- SIDEBAR CONTROL PANEL ---
st.sidebar.markdown("### ⚙️ Engine Control Panel")

# Check API Keys
gemini_key = os.getenv("GOOGLE_API_KEY")
groq_key = os.getenv("GROQ_API_KEY")

st.sidebar.markdown("**API Status Check:**")
if gemini_key:
    st.sidebar.markdown("🟢 Gemini API: Connected")
else:
    st.sidebar.markdown("🔴 Gemini API: Missing Key")
    
if groq_key:
    st.sidebar.markdown("🟢 Groq API: Connected")
else:
    st.sidebar.markdown("🔴 Groq API: Missing Key")

st.sidebar.markdown("---")

# Provider Selection
provider = st.sidebar.selectbox(
    "AI Provider",
    ["Gemini", "Groq"],
    index=0 if gemini_key else 1
)

# Model Selection
if provider == "Gemini":
    model_name = st.sidebar.selectbox(
        "Select Model",
        ["gemini-2.5-flash", "gemini-2.5-pro", "gemini-2.0-flash", "gemini-3.5-flash"],
        index=0
    )
else:
    model_name = st.sidebar.selectbox(
        "Select Model",
        ["llama-3.1-8b-instant", "llama-3.1-70b-versatile", "mixtral-8x7b-32768"],
        index=0
    )

temperature = st.sidebar.slider(
    "Temperature (Creativity)",
    min_value=0.0,
    max_value=1.0,
    value=0.2,
    step=0.1
)

st.markdown('<div class="breathing-header">🤖 InsightDocs AI Workspace</div>', unsafe_allow_html=True)
st.markdown('<div class="sub-title">Mobile-optimized deep document and image intelligence dashboard.</div>', unsafe_allow_html=True)

col1, col2 = st.columns(2)
with col1:
    uploaded_file = st.file_uploader("Choose Document File", type=["pdf", "docx", "txt", "md"], label_visibility="collapsed")
with col2:
    if st.button("🔄 Clear System Session/Cache", type="primary"):
        for key in list(st.session_state.keys()):
            del st.session_state[key]
        st.success("System memory and in-memory databases wiped clean!")
        st.rerun()

st.markdown("---")

if uploaded_file is not None:
    # Clear session instantly if file name changes to prevent data leak history
    if "current_file_name" in st.session_state:
        if st.session_state.current_file_name != uploaded_file.name:
            if "vectorstore" in st.session_state:
                st.session_state.vectorstore = None
                del st.session_state["vectorstore"]
            if "chat_history" in st.session_state:
                st.session_state.chat_history = []
                del st.session_state["chat_history"]
            st.session_state.current_file_name = uploaded_file.name
    else:
        st.session_state.current_file_name = uploaded_file.name

    if "vectorstore" not in st.session_state or st.session_state.vectorstore is None:
        with st.spinner("⚡ Processing complex document layout layers..."):
            text = extract_text_from_file(uploaded_file)
            if not text or not text.strip():
                st.error("❌ System Error: Unable to extract text data from this document layout. Please ensure the file contains valid digital text or images.")
                st.stop()
            chunks = split_text(text)
            st.session_state.vectorstore = create_vectorstore(chunks, get_embeddings())
        st.success(f"Loaded {len(chunks)} contextual nodes successfully via Isolated Local Deep parsing!")

    if "vectorstore" in st.session_state and st.session_state.vectorstore is not None:
        # Initialize chat history if not present
        if "chat_history" not in st.session_state:
            st.session_state.chat_history = []

        # Render past chat history in the visual stream
        for msg in st.session_state.chat_history:
            with st.chat_message(msg["role"]):
                st.markdown(msg["content"])
                if msg.get("sources"):
                    with st.expander("📚 Source Passages Used"):
                        for idx, src in enumerate(msg["sources"]):
                            st.markdown(f"**Source {idx+1}:**")
                            st.markdown(f'<div class="ai-response-box" style="font-size:0.9rem; padding:10px !important; margin-top:5px !important;">{src}</div>', unsafe_allow_html=True)

        # Chat Input Area at the bottom
        question = st.chat_input("💬 Ask anything about this document...")

        if question:
            # Instantly display user message in the visual stream
            with st.chat_message("user"):
                st.markdown(question)

            # Generate AI Response
            with st.spinner("Matching similarity nodes..."):
                try:
                    docs = st.session_state.vectorstore.similarity_search(question, k=6)
                    context = "\n\n".join([doc.page_content for doc in docs])

                    # Compile conversational memory
                    chat_history_str = ""
                    for msg in st.session_state.chat_history[-5:]:
                        chat_history_str += f"{msg['role'].upper()}: {msg['content']}\n"

                    prompt = f"""
You are a highly precise and knowledgeable Document Intelligence Assistant. Your objective is to answer the user's question, taking into account any previous conversation history and the provided document context.

Follow these instructions carefully:
1. **Analyze the Document Context**: Read the provided document context. Ground your answer in this document's details first.
2. **Supplement with General Knowledge**: If the document does not contain all the details needed to answer the question fully, or if the user asks for additional information/context, use your own pre-trained general knowledge to expand and explain.
3. **Clearly Distinguish Sources**: In your response, clearly state what information is extracted directly from the uploaded document, and what information is supplemented by your own general knowledge.
4. **Conversational Memory**: Use the chat history to understand follow-up references (e.g. "What about the second one?" or "Explain that").
5. Do NOT reuse general titles, watermarks, or first-page metadata expressions unless explicitly asked.

Chat History:
{chat_history_str}

Document Context:
{context}

Question:
{question}
"""
                    llm = get_llm(provider=provider, model_name=model_name, temperature=temperature)
                    response = llm.invoke(prompt)

                    # Display response in visual stream
                    with st.chat_message("assistant"):
                        st.markdown(response.content)
                        with st.expander("📚 Source Passages Used"):
                            for idx, doc in enumerate(docs):
                                st.markdown(f"**Source {idx+1}:**")
                                st.markdown(f'<div class="ai-response-box" style="font-size:0.9rem; padding:10px !important; margin-top:5px !important;">{doc.page_content}</div>', unsafe_allow_html=True)

                    # Save to session chat history
                    st.session_state.chat_history.append({"role": "user", "content": question})
                    st.session_state.chat_history.append({
                        "role": "assistant", 
                        "content": response.content,
                        "sources": [doc.page_content for doc in docs]
                    })
                    st.rerun()

                except Exception as e:
                    st.error(f"Inference Engine Error: {e}")
else:
    if "vectorstore" in st.session_state:
        del st.session_state["vectorstore"]
    if "current_file_name" in st.session_state:
        del st.session_state["current_file_name"]
    if "chat_history" in st.session_state:
        st.session_state.chat_history = []
        del st.session_state["chat_history"]
    st.info("💡 Tap 'Choose Document File' above to upload a document and activate the workspace.")